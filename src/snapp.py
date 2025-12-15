#!/usr/bin/env python3
import os
import sys
import webbrowser
import time
from pathlib import Path
from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
from typing import Optional, Tuple, List, Dict, Generator
import requests
import pandas as pd
from pandas.api.types import is_string_dtype
from bs4 import BeautifulSoup
import re
import string
import random
import argparse
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# =========================

# global constants and variables

FETCH_ONLY_MODE = False
OFFLINE_MODE = False
NORMAL_MODE = True
DEBUG_MODE = False

PUNCT = str.maketrans("", "", string.punctuation)

URL_RE = re.compile(r"https?://[^\s]+")
URL_TRAILING_PUNCT = ".,;:)]}>"  # stuff we often want to *not* include in the link

MAX_BLOCK_RETRIES_DEFAULT = 0
BLOCKING_SUSPECTED = False

NOT_FOUND_STRING = ""
NOT_FOUND_NAN = float('nan')

# =========================
# classes
# =========================

class GSBlockedError(Exception):
    """Raised when those Google yahoos appear to be blocking requests."""
    pass

class AuthorMatchError(RuntimeError):
    """Raised when no authors or multiple in list match the candidate's Google Scholar name.
    This is treated as fatal.
    """
    pass

class AuthorMatchWarning(RuntimeError):
    """Raised when no authors or multiple in list match the candidate's Google Scholar name.
    This is treated as a warning.
    """
    pass

# =========================
# helper functions
# =========================

# extract user_id from URL

def user_id_from_url(url: str) -> Optional[str]:

    parsed = urlparse(url)
    qs = parse_qs(parsed.query)
    user_vals = qs.get("user")
    if user_vals:
        return user_vals[0]
    return None

# =========================

# sanitise list of URLs to standard format in English!

def sanitise_urls(urls: List[str]) -> List[str]:

    sanitised: List[str] = []
    for url in urls:
        user_id = user_id_from_url(url)
        if user_id:
            sanitised_url = f"https://scholar.google.com/citations?user={user_id}&hl=en"
            sanitised.append(sanitised_url)
        else:
            print(f" Warning - Could not extract user id from URL: {url}, skipping.")
    return sanitised

# =========================

# sanitise a single URL to standard format in English!

def sanitise_url(url: str) -> Optional[str]:
    user_id = user_id_from_url(url)
    if user_id:
        sanitised_url = f"https://scholar.google.com/citations?user={user_id}&hl=en"
        return sanitised_url
    else:
        print(f" Warning - Could not extract user id from URL: {url}. sanitise_url failed!")
        return None

# ========================

# normalise a journal name by cleaning punctuation and whitespace

def normalise_journal_name(name: str) -> str:
    name = name.lower().strip()
    # remove punctuation but keep spaces
    name = name.translate(PUNCT)
    # collapse multiple spaces
    name = re.sub(r"\s+", " ", name)
    return name

# ========================

# extract the journal name from the journal_info field

def extract_journal_name(raw_info: str) -> str:

    s = raw_info.strip()

    # look for the first spot where a volume/pages/year chunk starts.
    # this is usually: space + digit, or comma + space + digit, or space + '(' + digit
    m = re.match(r"^(.*?)(?=\s\d|\s\(\d|,\s*\d{1,4})", s)
    if m:
        return m.group(1).strip()
    return s  # fallback: whole string if no match

# ========================

# compare author name to candidate name
# candidate name is assumed to be full name format "Firstname Middlename Surname"
# author name is assumed to be in "Initials Surname" format

def compare_initialled_name_with_full_name(
    initialled_name: str,
    full_name: str
) -> bool:
    
    global DEBUG_MODE
    
    if DEBUG_MODE: print(f"\n Comparing initialled name '{initialled_name}' with full name '{full_name}'")
    
    # search for a hyphen in the intialled name to detect multi-barrelled surnames
    found_hyphen_initialled_name = False
    if "-" in initialled_name:
        found_hyphen_initialled_name = True
        if DEBUG_MODE: print(f"   Detected hyphen in initialled name '{initialled_name}' => multi-barrelled surname assumed.")
        # count the number of hyphens in the full name
        hyphen_count = initialled_name.count("-")
        if hyphen_count > 1:
            print(f"   ERROR - Multiple hyphens ({hyphen_count}) detected in initialled name '{initialled_name}'")
            exit(1)        
        
    found_hyphen_full_name = False
    if "-" in full_name:
        found_hyphen_full_name = True
        if DEBUG_MODE: print(f"   Detected hyphen in full name '{full_name}' => multi-barrelled surname assumed.")
        # count the number of hyphens in the full name
        hyphen_count = full_name.count("-")
        if hyphen_count > 1:
            print(f"   ERROR - Multiple hyphens ({hyphen_count}) detected in full name '{full_name}'")
            exit(1)
    
    # decompose full name and reconstruct as an initialled name
    full_name_parts = full_name.strip().split(" ")    
    
    if len(full_name_parts) >= 2:
        full_name_initials = ""
        # initial assumption that this is not a multi-barrelled surname
        if DEBUG_MODE: print(f"   Last word -> forms surname base: {full_name_parts[-1]}")
        full_name_surname = full_name_parts[-1]
        # add any preceding parts that start with lowercase letters to the surname otherwise treat as initials
        for part in reversed(full_name_parts[:-1]):
            if not part.isalpha():
                if DEBUG_MODE: print(f"   '{part}' contains non-alphabetic characters, skipping.")
                continue
            elif part[0] != part[0].upper():
                if DEBUG_MODE: print(f"   First letter '{part[0]}' is not uppercase => add to surname.")
                full_name_surname = part + " " + full_name_surname
            else:
                if DEBUG_MODE: print(f"   First letter '{part[0]}' is uppercase => add it to the initials string.")
                full_name_initials = part[0] + full_name_initials
    else:
        full_name_initials = ""
        full_name_surname = full_name_parts[0]       
    
    full_name_initialised = " ".join([full_name_initials, full_name_surname]) 
    
    # remove hyphens from both names if either has a hyphen
    if found_hyphen_initialled_name or found_hyphen_full_name:
        initialled_name = initialled_name.replace("-", " ")
        full_name_initialised = full_name_initialised.replace("-", " ")
        if DEBUG_MODE: print(f"   Hyphens removed for comparison: '{initialled_name}' and '{full_name_initialised}'")
    
    # compare names
    if DEBUG_MODE: print(f"   Comparing names: '{initialled_name}' with '{full_name_initialised}'")
    
    if initialled_name.lower() != full_name_initialised.lower():
        if DEBUG_MODE: print(f"   Return False <= Names do not match: '{initialled_name}' != '{full_name_initialised}")
        return False
    
    if DEBUG_MODE: print(f"  Return True <=  '{initialled_name}' == '{full_name_initialised}'")
    return True

# =========================

# random sleep
        
def random_sleep(typical_delay: float) -> None:

        sleep_s = random.uniform(typical_delay * 0.5, typical_delay * 1.5)
        print(f" Random (human-like) delay for {sleep_s:.1f} seconds before proceeding...")
        time.sleep(sleep_s)
        
# =========================

# deal with GS blocking, CAPTCHA and other antics

def looks_like_block_page(html: str) -> bool:

    text = html.lower()
    block_markers = [
        "unusual traffic",
        "/sorry/",
        "not a robot",
        "solve the captcha",
        "submit a verification",
        "our systems have detected",
        "scholar help",
    ]

    if "captcha" in text:
        return True

    return any(marker in text for marker in block_markers)

# ===================

# get the list_works URL using cstart/pagesize params

def build_list_works_url(base_url: str, cstart: int, pagesize: int = 100) -> str:

    parsed = urlparse(base_url)
    qs = parse_qs(parsed.query)

    # make sure key params are present; add/override view_op, cstart, pagesize
    qs["view_op"] = ["list_works"]
    qs["cstart"] = [str(cstart)]
    qs["pagesize"] = [str(pagesize)]

    new_query = urlencode(qs, doseq=True)
    new_parsed = parsed._replace(query=new_query)
    return urlunparse(new_parsed)

# =========================

# step through pages with requests

def iter_scholar_pages_requests(
    base_url: str,
    session: requests.Session,
    pagesize: int = 100,
    max_pages: int = 50,
    delay: float = 8.0,                                     # typical delay between successful pages
    max_block_retries: int = MAX_BLOCK_RETRIES_DEFAULT,     # how many times to retry a blocked page
    block_backoff_base: float = 10.0                        # starting backoff in seconds
) -> Generator[str, None, None]:

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0 Safari/537.36"
        )
    }

    cstart = 0
    page_index = 0

    while page_index < max_pages:
        url = build_list_works_url(base_url, cstart=cstart, pagesize=pagesize)
        print(f"\n Loading publications page {page_index + 1} (cstart={cstart})")

        block_attempts = 0

        while True:
            try:
                resp = session.get(url, headers=headers, timeout=15)
            except requests.RequestException as e:
                print(f"  Error: request exception {type(e).__name__}: {e}")
                block_attempts += 1
                if block_attempts > max_block_retries:
                    print("  Too many request failures for this page, giving up.")
                    return
                backoff = block_backoff_base * (2 ** (block_attempts - 1))
                print(f"  Backing off for {backoff:.1f} seconds before retrying...")
                random_sleep(backoff)
                continue

            # deal with status-based blocking
            if resp.status_code in (429, 503):
                print(f"  HTTP {resp.status_code} suggests rate limiting or temporary block.")
                print("  Stopping pagination for this profile.")
                return

            if resp.status_code != 200:
                print(f"  Error: HTTP {resp.status_code}, stopping.")
                return

            html = resp.text

            # check for CAPTCHA / unusual traffic page
            if looks_like_block_page(html):
                print("  Page looks like a CAPTCHA / 'unusual traffic' block.")
                print("  Stopping pagination for this profile.")
                return

            # woohoo - we have a page
            break

        # sanity check - parse the publications table to see if we have rows
        soup = BeautifulSoup(html, "html.parser")
        table_pubs = soup.find("table", id="gsc_a_t")
        if not table_pubs:
            print("  No publications table found, stopping.")
            return

        rows = table_pubs.find_all("tr", class_="gsc_a_tr")
        if not rows:
            print("  No publication rows found, stopping.")
            return

        print(f"  Found {len(rows)} publication rows on this page.")
        yield html

        # if we have fewer rows than pagesize then this is the last page
        if len(rows) < pagesize:
            print("  Last page detected (fewer than pagesize rows).")
            return

        cstart += pagesize
        page_index += 1

        # add delay between pages to avoid looking like a bot 
        random_sleep(delay)

# =========================

# I love BeautifulSoup :)~

def scrape_it(
    html: str,
    journal_list: List[str],
    normalised_journal_titles: Dict[str, str],
    page_idx: int,
) -> Tuple[
    Optional[str],          # name
    Optional[str],          # institution
    Optional[List[str]],    # research areas
    Optional[int],          # h_all
    Optional[int],          # h_5y
    Optional[int],          # cit_all
    Optional[int],          # cit_5y
    Dict[str, int],         # journal_match_counts
    Dict[str, int],         # journal_match_counts_fa
    Dict[str, int],         # journal_match_counts_sa
    Dict[str, int],         # journal_match_counts_la
    Dict[str, int],         # journal_num_authors
    Dict[str, List[str]],   # journal_match_details
    int,                    # article_count
]:
    global FETCH_ONLY_MODE
    global DEBUG_MODE
    
    soup = BeautifulSoup(html, "html.parser")

    # ---------------------------------------------------------------------
    # name tag: <div id="gsc_prf_in">Name</div>
    # ---------------------------------------------------------------------
    candidate_gs_name: Optional[str] = None
    name_div = soup.find("div", id="gsc_prf_in")
    if name_div:
        candidate_gs_name = name_div.get_text(strip=True)

    if candidate_gs_name:
        print(f"\n Scraping profile page {page_idx + 1} for {candidate_gs_name}")
    else:
        print(f"\n Scraping profile page {page_idx + 1} for 'UNKNOWN'")

    # defaults for front matter 
    institution: Optional[str] = None
    research_areas: Optional[List[str]] = None
    h_all: Optional[int] = None
    h_5y: Optional[int] = None
    cit_all: Optional[int] = None
    cit_5y: Optional[int] = None

    # get the front matter data only on the first page
    # ---------------------------------------------------------------------
    # institution / affiliation tag:
    #   <div class="gsc_prf_il">The University of Excellence and other Buzzwords</div>
    # ---------------------------------------------------------------------
    if page_idx == 0:
        # institution
        inst_divs = soup.find_all("div", class_="gsc_prf_il")
        if inst_divs:
            institution = inst_divs[0].get_text(strip=True)
        if DEBUG_MODE: print(f"\n Institution: {institution if institution else 'None found'}")

        # ---------------------------------------------------------------------
        # research areas / interests tags:
        #   <div id="gsc_prf_int">
        #       <a class="gsc_prf_inta">Area 1</a>
        #       <a class="gsc_prf_inta">Area 2</a>
        #   </div>
        # ---------------------------------------------------------------------
        ra: List[str] = []
        int_div = soup.find("div", id="gsc_prf_int")
        if int_div:
            for a in int_div.find_all("a", class_="gsc_prf_inta"):
                text = a.get_text(strip=True)
                if text:
                    ra.append(text)
        research_areas = ra

        if DEBUG_MODE:
            if research_areas:
                print(" Research areas: " + ", ".join(research_areas))
            else:
                print(" Research areas: None found")

        # ---------------------------------------------------------------------
        # h-index and citations table tags:
        # <table id="gsc_rsb_st">
        #   rows for "Citations", "h-index", "i10-index"
        #   columns: [label, All, Since YYYY]
        # ---------------------------------------------------------------------
        table = soup.find("table", id="gsc_rsb_st")
        if table:
            for row in table.find_all("tr"):
                cells = row.find_all("td")
                if not cells:
                    continue

                label = cells[0].get_text(strip=True).lower()

                if "citations" in label:
                    if len(cells) >= 2:
                        try:
                            cit_all = int(cells[1].get_text(strip=True))
                        except ValueError:
                            cit_all = None
                    if len(cells) >= 3:
                        try:
                            cit_5y = int(cells[2].get_text(strip=True))
                        except ValueError:
                            cit_5y = None

                elif "h-index" in label:
                    if len(cells) >= 2:
                        try:
                            h_all = int(cells[1].get_text(strip=True))
                        except ValueError:
                            h_all = None
                    if len(cells) >= 3:
                        try:
                            h_5y = int(cells[2].get_text(strip=True))
                        except ValueError:
                            h_5y = None
        if DEBUG_MODE:
            print(f" h-index (all): {h_all}, h-index (5y): {h_5y}")
            print(f" citations (all): {cit_all}, citations (5y): {cit_5y}")

    # journal matching on all pages
    # ---------------------------------------------------------------------
    # journal matching in publications table
    # <table id="gsc_a_t">...</table>
    # ---------------------------------------------------------------------
    journal_match_counts: Dict[str, int] = {j: 0 for j in journal_list}
    journal_match_counts_fa: Dict[str, int] = {j: 0 for j in journal_list}
    journal_match_counts_sa: Dict[str, int] = {j: 0 for j in journal_list}
    journal_match_counts_la: Dict[str, int] = {j: 0 for j in journal_list}
    journal_num_authors: Dict[str, int] = {j: 0 for j in journal_list}
    journal_match_details: Dict[str, List[str]] = {j: [] for j in journal_list}
    article_count = 0

    table_pubs = soup.find("table", id="gsc_a_t")
    if table_pubs and journal_list:
        for row in table_pubs.find_all("tr", class_="gsc_a_tr"):
            cells = row.find_all("td", class_="gsc_a_t")
            if not cells:
                continue

            td = cells[0]
            gray_elems = td.find_all("div", class_="gs_gray")
            # expect at least 2 gs_gray divs:
            # [0] authors
            # [1] journal info
            if len(gray_elems) < 2:
                continue

            article_count += 1

            raw_info = gray_elems[1].get_text(strip=True)

            # extract the journal title from the raw info string
            journal_title = extract_journal_name(raw_info)

            # remove punctuation and normalise
            journal_norm = normalise_journal_name(journal_title)

            # compare against normalised journal titles list
            matched_journal = normalised_journal_titles.get(journal_norm)

            if matched_journal:
                print(f" Journal match: '{raw_info}' -> '{matched_journal}'")
                journal_match_counts[matched_journal] += 1

                # ---- capture full publication details ----
                title_elem = td.find("a", class_="gsc_a_at")
                title = title_elem.get_text(strip=True) if title_elem else "UNKNOWN TITLE"

                authors = gray_elems[0].get_text(strip=True)
                journal_info = gray_elems[1].get_text(strip=True)
                
                # ---- cited-by + year live in sibling columns ----
                cited_by = 0
                year = ""

                cited_td = row.find("td", class_="gsc_a_c")
                if cited_td:
                    cited_a = cited_td.find("a")  # when citations exist it's usually a link
                    cited_txt = (cited_a.get_text(strip=True) if cited_a else cited_td.get_text(strip=True))
                    try:
                        cited_by = int(cited_txt) if cited_txt else 0
                    except ValueError:
                        cited_by = 0

                cited_by_url = ""
                if cited_td:
                    cited_a = cited_td.find("a")
                    if cited_a and cited_a.get("href"):
                        cited_by_url = cited_a["href"]

                year_td = row.find("td", class_="gsc_a_y")
                if year_td:
                    year_txt = year_td.get_text(strip=True)
                    year = year_txt  # keep as string; you can int() it if you want

                # create a list of authors by separating on commas
                author_list = [a.strip() for a in authors.split(",") if a.strip()]
                
                if DEBUG_MODE: print(f" Author list: {author_list}")
                                
                # determine which authors match the candidate's name
                highlighted_author_list = []
                counter = 0
                count_highlighted = 0
                for a in author_list:
                    counter += 1
                    
                    if compare_initialled_name_with_full_name(
                        initialled_name = a,
                        full_name = candidate_gs_name or ""
                    ):
                        if DEBUG_MODE: print(f"  Matched author: {a} with candidate's name {candidate_gs_name}")
                        # highlight the matched author
                        highlighted_author_list.append(f"**{a}**")
                        count_highlighted += 1
                        if counter == 1:
                            journal_match_counts_fa[matched_journal] += 1
                        elif counter == 2:
                            journal_match_counts_sa[matched_journal] += 1
                        elif counter == len(author_list):
                            journal_match_counts_la[matched_journal] += 1
                    else:
                        if DEBUG_MODE: print(f"  Did not match author: {a} with candidate name {candidate_gs_name}")
                        highlighted_author_list.append(a)
                                    
                '''if count_highlighted == 0:
                    raise AuthorMatchWarning(
                        f"No author matched candidate '{candidate_gs_name}' "
                        f"for journal '{matched_journal}'. "
                        f"Authors found: {', '.join(author_list)}"
                    )
                elif count_highlighted > 1:
                    raise AuthorMatchWarning(
                        f"Multiple ({count_highlighted}) authors matched candidate "
                        f"'{candidate_gs_name}' for journal '{matched_journal}'. "
                        f"Authors found: {', '.join(author_list)}"
                    )'''
                    
                journal_num_authors[matched_journal] += len(author_list)
                
                authors = ", ".join(highlighted_author_list)

                full_entry = f'{authors} | {title} | {journal_info} | {cited_by} | {year}'
                journal_match_details[matched_journal].append(full_entry)

    if not FETCH_ONLY_MODE and DEBUG_MODE:
        print(f" Total articles on this page: {article_count}")
        print(" Journal match counts on this page:")
        for j, c in journal_match_counts.items():
            if c > 0:
                print(f"  {j}: {c}")

    print("\n ------------------------------------------\n")

    return (
        candidate_gs_name,
        institution,
        research_areas,
        h_all,
        h_5y,
        cit_all,
        cit_5y,
        journal_match_counts,
        journal_match_counts_fa,
        journal_match_counts_sa,
        journal_match_counts_la,
        journal_num_authors,
        journal_match_details,
        article_count,
    )
            
# ========================

# step through all GS pages and scrape profile info 

def scrape_profile_all_publications(
    profile_url: str,
    journal_list: List[str],
    normalised_journal_titles: Dict[str, str],
    html_dir: str = "./html",
    max_pages: int = 50,
) -> Tuple[
    Optional[str],          # name
    Optional[str],          # institution
    List[str],              # research_areas
    Optional[int],          # h_all
    Optional[int],          # h_5y
    Optional[int],          # cit_all
    Optional[int],          # cit_5y
    Dict[str, int],         # total_journal_match_counts
    Dict[str, int],         # total_journal_match_counts_fa
    Dict[str, int],         # total_journal_match_counts_sa
    Dict[str, int],         # total_journal_match_counts_la
    Dict[str, int],         # total_journal_num_authors
    Dict[str, List[str]],   # total_journal_details
    int,                    # total_article_count
    bool                    # any_page
]:

    name: Optional[str] = None
    institution: Optional[str] = None
    research_areas: List[str] = []
    h_all: Optional[int] = None
    h_5y: Optional[int] = None
    cit_all: Optional[int] = None
    cit_5y: Optional[int] = None

    total_journal_counts: Dict[str, int] = {j: 0 for j in journal_list}
    total_journal_counts_fa: Dict[str, int] = {j: 0 for j in journal_list}
    total_journal_counts_sa: Dict[str, int] = {j: 0 for j in journal_list}
    total_journal_counts_la: Dict[str, int] = {j: 0 for j in journal_list}
    total_journal_num_authors: Dict[str, int] = {j: 0 for j in journal_list}
    total_journal_details: Dict[str, List[str]] = {j: [] for j in journal_list}
    total_article_count = 0

    any_page = False
    user_id = user_id_from_url(profile_url) or "UNKNOWN"

    base_dir = Path(html_dir)
    page_idx = 0

    for page_num in range(1, max_pages + 1):
        path = base_dir / f"{user_id}_p{page_num}.htm"
        if not path.exists():
            # no more cached pages
            break

        any_page = True

        print(f" Loading cached HTML for {user_id} page {page_num} -> {path}")

        with open(path, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()

        # scrape the page
        (
            name,
            institution,
            research_areas,
            h_all,
            h_5y,
            cit_all,
            cit_5y,
            page_journal_counts,
            page_journal_counts_fa,
            page_journal_counts_sa,
            page_journal_counts_la,
            page_journal_num_authors,
            page_journal_details,
            page_article_count,
        ) = scrape_it(html, journal_list, normalised_journal_titles, page_idx)   

        # accumulate journal counts, details and article counts
        for j in journal_list:
            total_journal_counts[j] += page_journal_counts.get(j, 0)
            total_journal_counts_fa[j] += page_journal_counts_fa.get(j, 0)
            total_journal_counts_sa[j] += page_journal_counts_sa.get(j, 0)
            total_journal_counts_la[j] += page_journal_counts_la.get(j, 0)
            total_journal_num_authors[j] += page_journal_num_authors.get(j, 0)
            total_journal_details[j].extend(page_journal_details[j])

        total_article_count += page_article_count
        
        page_idx += 1

    if not any_page:
        print(f"\n Warning - No cached HTML pages found for user_id={user_id} in {html_dir}\n")
    else:
        if DEBUG_MODE:
            print(" === Results aggregated over ALL cached pages ===\n")
            print(f" Name: {name}")
            print(f" Institution: {institution}")
            if research_areas:
                print(" Research areas: " + ", ".join(research_areas))
            print(f" h-index (all): {h_all}, h-index (5y): {h_5y}")
            print(f" citations (all): {cit_all}, citations (5y): {cit_5y}")
            print(f" Total articles (all pages): {total_article_count}")
            print(" Journal match counts (all pages):")
            for journal, count in total_journal_counts.items():
                if count > 0:
                    print(f"  {journal}: {count}")
            print(" First author journal match counts (all pages):")                
            for journal, count in total_journal_counts_fa.items():
                if count > 0:
                    print(f"  {journal}: {count}")     
            print(" Second author journal match counts (all pages):")
            for journal, count in total_journal_counts_sa.items():
                if count > 0:
                    print(f"  {journal}: {count}")
            print(" Last author journal match counts (all pages):")
            for journal, count in total_journal_counts_la.items():
                if count > 0:
                    print(f"  {journal}: {count}")           
            print("\n ===============================================================================\n")

    return (
        name,
        institution,
        research_areas,
        h_all,
        h_5y,
        cit_all,
        cit_5y,
        total_journal_counts,
        total_journal_counts_fa,
        total_journal_counts_sa,
        total_journal_counts_la,
        total_journal_num_authors,
        total_journal_details,
        total_article_count,
        any_page
    )

# =========================

# fetch and cache profile HTML only

def fetch_and_cache_profile(
    candidate: tuple,
    session: requests.Session,
    pagesize: int = 100,
    max_pages: int = 50,
    delay: float = 8.0,    
    max_block_retries: int = MAX_BLOCK_RETRIES_DEFAULT,
    block_backoff_base: float = 10.0,
    html_dir: str = "./html",
) -> bool:

    
    print(f" === Fetch and cache: candidate {candidate.candidate_id}: {candidate.candidate_name} ===")

    url = candidate.gs_url

    if pd.isna(url):
        print("\n  Warning - Empty Google Scholar Link, skipping profile.")
        return False

    url_str = str(url).strip()
    print(f"\n  Attempting to fetch and cache Google Scholar URL: {url_str}")

    sanitised_url = sanitise_url(url_str)
    if sanitised_url is None:
        print("\n  Warning - Could not sanitise URL, skipping profile.")
        return False

    any_page = False
    user_id = user_id_from_url(sanitised_url) or "UNKNOWN"

    # set up cache directory
    out_dir = Path(html_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    for page_idx, html in enumerate(
        iter_scholar_pages_requests(
            sanitised_url,
            session,
            pagesize=pagesize,
            max_pages=max_pages,
            delay=delay,
            max_block_retries=max_block_retries,
            block_backoff_base=block_backoff_base,
        )
    ):
        any_page = True
        
        page_num = page_idx + 1
        out_path = Path(html_dir) / f"{user_id}_p{page_num}.htm"
        with open(out_path, "w", encoding="utf-8", errors="replace") as f:
            f.write(html)
        print(f"  Cached HTML for {user_id} page {page_num} -> {out_path}")

    if not any_page:
        print(f"\n Warning - No publication pages cached for URL: {sanitised_url}")
        # likely blocked or unreachable
        raise GSBlockedError(f"Blocked or no pages for {sanitised_url}")
        
    return True

# =========================

# create a single string object that gives a summary of the candidate record and journal_details   
 
def create_summary(
    record: Dict[str, object],
    journal_counts: Dict[str, int] = {},
    journal_counts_fa: Dict[str, int] = {},
    journal_counts_sa: Dict[str, int] = {},
    journal_counts_la: Dict[str, int] = {},
    journal_num_authors: Dict[str, int] = {},    
    journal_details: Dict[str, List[str]] = {},
    journal_list: List[str] = [],
    is_empty_record: bool = False,
    markdown: bool = True # otherwise plain text
) -> str:
    
    global DEBUG_MODE
    
    summary_lines: List[str] = []
    
    if markdown:
        summary_lines.append(f"# Candidate: {record.get('candidate_id', 'UNKNOWN')} - {record.get('candidate_name', 'UNKNOWN')}")
        summary_lines.append(f"**Gender**: {record.get('gender', 'UNKNOWN')}")
        summary_lines.append(f"**Country of residence**: {record.get('country', 'UNKNOWN')}")
        #summary_lines.append(f"Email: {record.get('email', 'UNKNOWN')}")
        summary_lines.append(f"**Current Employee**: {record.get('current_employee', 'UNKNOWN')}")
        summary_lines.append(f"**Expertise Area**: {record.get('expertise_area', 'UNKNOWN')}")
        summary_lines.append(f"**Academic Level**: {record.get('academic_level', 'UNKNOWN')}")
        summary_lines.append(f"**PhD Year**: {record.get('PhD_year', 'UNKNOWN')}")
        summary_lines.append(f"**PhD Institution**: {record.get('PhD_institution', 'UNKNOWN')}")
        summary_lines.append(f"**PhD Institution Rank**: {record.get('PhD_institution_rank', 'UNKNOWN')}") 
    else:
        summary_lines.append(f"Candidate: {record.get('candidate_id', 'UNKNOWN')} - {record.get('candidate_name', 'UNKNOWN')}")
        summary_lines.append(f"Gender: {record.get('gender', 'UNKNOWN')}")
        summary_lines.append(f"Country of residence: {record.get('country', 'UNKNOWN')}")
        #summary_lines.append(f"Email: {record.get('email', 'UNKNOWN')}")
        summary_lines.append(f"Current Employee: {record.get('current_employee', 'UNKNOWN')}")
        summary_lines.append(f"Expertise Area: {record.get('expertise_area', 'UNKNOWN')}")
        summary_lines.append(f"Academic Level: {record.get('academic_level', 'UNKNOWN')}")
        summary_lines.append(f"PhD Year: {record.get('PhD_year', 'UNKNOWN')}")
        summary_lines.append(f"PhD Institution: {record.get('PhD_institution', 'UNKNOWN')}")
        summary_lines.append(f"PhD Institution Rank: {record.get('PhD_institution_rank', 'UNKNOWN')}")
        summary_lines.append("")    
        summary_lines.append("------------------------------------------")
    
    if is_empty_record:
        summary_lines.append("")
        if markdown:
            summary_lines.append("### No Google Scholar profile data found.")
        else:
            summary_lines.append("No Google Scholar profile data found.")
        summary_lines.append("")
        summary = "\n".join(summary_lines)
        
        if True:
            if markdown:
                print("\n ======= FULL SUMMARY - Markdown ======= \n")
            else:
                print("\n ======= FULL SUMMARY ======= \n")
            print(summary)
            print("\n ============================ \n")
        
        return summary
    
    summary_lines.append("")
    if markdown:
        summary_lines.append(f"## Google Scholar Profile Summary")
        summary_lines.append(f"**URL**: {record.get('gs_url', 'UNKNOWN')}")
        summary_lines.append("")
        summary_lines.append(f"**Current Institution**: {record.get('gs_institution', 'UNKNOWN')}")
        summary_lines.append(f"**Research Areas**: {record.get('gs_research_areas', 'UNKNOWN')}")
        summary_lines.append(f"**Citations (All | 5y)**: {record.get('citations_all', 0)} | {record.get('citations_5y', 0)}")
        summary_lines.append(f"**h-index (All | 5y)**: {record.get('h_index_all', 0)} | {record.get('h_index_5y', 0)}")
        summary_lines.append(f"**Total Articles**: {record.get('article_count', 0)}")
        summary_lines.append(f"**Total Articles in the Journal List**: {record.get('journal_count_tot', 0)}")
        summary_lines.append(f"**Total First Author Papers**: {record.get('journal_count_tot_fa', 0)}")
        summary_lines.append(f"**Total Second Author Papers**: {record.get('journal_count_tot_sa', 0)}")
        summary_lines.append(f"**Total Last Author Papers**: {record.get('journal_count_tot_la', 0)}")
        #average_num_authors = record.get('journal_average_num_authors', 0)
        #summary_lines.append(f"Average Number of Authors per Paper: {average_num_authors:.1f}")
        summary_lines.append("")
    else:
        summary_lines.append(f"Google Scholar Profile Summary:")
        #summary_lines.append(f"URL: {record.get('gs_url', 'UNKNOWN')}")
        summary_lines.append(f"Current Institution: {record.get('gs_institution', 'UNKNOWN')}")
        summary_lines.append(f"Research Areas: {record.get('gs_research_areas', 'UNKNOWN')}")
        summary_lines.append(f"Citations (All | 5y): {record.get('citations_all', 0)} | {record.get('citations_5y', 0)}")
        summary_lines.append(f"h-index (All | 5y): {record.get('h_index_all', 0)} | {record.get('h_index_5y', 0)}")
        summary_lines.append("")
        summary_lines.append(f"Total Articles: {record.get('article_count', 0)}")
        summary_lines.append(f"Total Articles in the Journal List: {record.get('journal_count_tot', 0)}")
        summary_lines.append(f"Total First Author Papers: {record.get('journal_count_tot_fa', 0)}")
        summary_lines.append(f"Total Second Author Papers: {record.get('journal_count_tot_sa', 0)}")
        summary_lines.append(f"Total Last Author Papers: {record.get('journal_count_tot_la', 0)}")
        #average_num_authors = record.get('journal_average_num_authors', 0)
        #summary_lines.append(f"Average Number of Authors per Paper: {average_num_authors:.1f}")
        summary_lines.append("")
        summary_lines.append("------------------------------------------")
        summary_lines.append("")        
      
    if markdown:    
        summary_lines.append("## Journal / Conference Article Summary")
    else:
        summary_lines.append("Journal / Conference Article Summary:")
        summary_lines.append("(Articles from Journal List Only)")
    
    for journal in journal_list:
        count = journal_counts.get(journal, 0)
        count_fa = journal_counts_fa.get(journal, 0)
        count_sa = journal_counts_sa.get(journal, 0)
        count_la = journal_counts_la.get(journal, 0)
        num_authors = journal_num_authors.get(journal, 0)
        details = journal_details.get(journal, [])

        if count > 0:
            if markdown:
                summary_lines.append(f"### {journal}")
                summary_lines.append(f"**Number of articles**: {count}")   
                summary_lines.append(f"**Number of first author articles**: {count_fa}")   
                summary_lines.append(f"**Number of second author articles**: {count_sa}")   
                summary_lines.append(f"**Number of last author articles**: {count_la}")
            else:
                summary_lines.append(f"{journal}")
                summary_lines.append(f"Number of articles: {count}")   
                summary_lines.append(f"Number of first author articles: {count_fa}")   
                summary_lines.append(f"Number of second author articles: {count_sa}")   
                summary_lines.append(f"Number of last author articles: {count_la}")                

            #summary_lines.append(f"Average number of authors: {num_authors / count:.1f}")

            for detail in details:
                # expected detail format: authors | title | journal_info
                parts = [p.strip() for p in detail.split("|", 4)]

                if markdown:
                    if len(parts) == 5:
                        authors, title, journal_info, cited_by, year = parts
                        # replace cited_by= with just the number
                        # cited_by = cited_by.replace("cited_by=", "cited by ").strip()
                        line = f'- {authors}, "{title}", {journal_info} **[{cited_by}]**'
                    else:
                        # fallback if format is unexpected
                        line = '-' + detail.replace("|", ", ")

                    summary_lines.append(line)
                else:   
                    if len(parts) == 5:
                        authors, title, journal_info, cited_by, year = parts
                        # replace cited_by= with just the number
                        # cited_by = cited_by.replace("cited_by=", "cited by ").strip()
                        line = f'- {authors}, "{title}", {journal_info} [{cited_by}]'
                    else:
                        # fallback if format is unexpected
                        line = '-' + detail.replace("|", ", ")

                    summary_lines.append(line)
                        

            # blank line between journals for readability
            summary_lines.append("")

    # remove trailing line
    while summary_lines and summary_lines[-1] == "":
        summary_lines.pop()
                
    summary = "\n".join(summary_lines)
    
    # replace all instances in which there is no space after a comma with a space
    # or no space around parentheses
    summary = normalise_punctuation(summary)
    
    if True:
        if markdown:
            print("\n ======= FULL SUMMARY - Markdown ======= \n")
        else:
            print("\n ======= FULL SUMMARY ======= \n")
        print(summary)
        print("\n ============================ \n")
        
    return summary
   
# ========================

def normalise_punctuation(summary: str) -> str:
    # ensure space after comma
    summary = re.sub(r",(?=\S)", ", ", summary)

    # ensure space before opening parenthesis
    summary = re.sub(r"(?<![ \n])\(", " (", summary)

    # ensure space after closing parenthesis
    summary = re.sub(r"\)(?=\S)", ") ", summary)

    # collapse multiple SPACES only (not newlines)
    summary = re.sub(r"[ ]{2,}", " ", summary)

    return summary.strip()

# =========================

# Add a clickable hyperlink to a paragraph.

def add_hyperlink(paragraph, url: str, text: str | None = None):

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # make it look like a typical hyperlink (blue + underline).
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text if text is not None else url
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# =======================

# peel trailing punctuation off the URL (so the link is clean)

def _split_url_trailing_punct(url: str) -> tuple[str, str]:
    
    trail = ""
    while url and url[-1] in URL_TRAILING_PUNCT:
        trail = url[-1] + trail
        url = url[:-1]
    return url, trail

# =======================

# adds runs to an existing paragraph
def add_md_inline_runs(p, text: str) -> None:
    i = 0
    n = len(text)

    while i < n:
        # find next URL / bold / italic marker
        m_url = URL_RE.search(text, i)
        j_url = m_url.start() if m_url else -1

        j_bold = text.find("**", i)
        j_ital = text.find("_", i)

        candidates = [(j_url, "url"), (j_bold, "bold"), (j_ital, "ital")]
        candidates = [(pos, kind) for pos, kind in candidates if pos != -1]
        if not candidates:
            p.add_run(text[i:])
            return

        j, kind = min(candidates, key=lambda x: x[0])

        # emit normal text before the token
        if j > i:
            p.add_run(text[i:j])

        if kind == "url":
            raw_url = m_url.group(0)
            url, trailing = _split_url_trailing_punct(raw_url)

            # hyperlink run
            add_hyperlink(p, url, text=url)

            # trailing punctuation after the link (normal text)
            if trailing:
                p.add_run(trailing)

            i = m_url.end()
            continue

        if kind == "bold":
            k = text.find("**", j + 2)
            if k == -1:
                # unmatched -> literal
                p.add_run(text[j:])
                return
            run_text = text[j + 2 : k]
            r = p.add_run(run_text)
            r.bold = True
            i = k + 2
            continue

        # ital
        k = text.find("_", j + 1)
        if k == -1:
            p.add_run(text[j:])
            return
        run_text = text[j + 1 : k]
        r = p.add_run(run_text)
        r.italic = True
        i = k + 1

# ========================

# adds a single markdown-ish line to a docx document

def add_md_line(doc: "DocxDocument", line: str) -> None:

    raw = line.rstrip("\n")
    if not raw.strip():
        doc.add_paragraph("")  # blank line
        return

    # headings: expects '# Title" with a space after hashes
    m = re.match(r"^(#{1,6})\s+(.*)$", raw)
    if m:
        level = min(len(m.group(1)), 6)
        text = m.group(2).strip()
        p = doc.add_paragraph(style=f"Heading {level}")
        add_md_inline_runs(p, text)
        return

    # bullets
    if raw.lstrip().startswith("- "):
        text = raw.lstrip()[2:].strip()
        p = doc.add_paragraph(style="List Bullet")
        add_md_inline_runs(p, text)
        return

    # normal paragraph
    p = doc.add_paragraph()
    add_md_inline_runs(p, raw)

# =========================

def set_style_font(style, name="Arial", size_pt=10):
    font = style.font
    font.name = name
    font.size = Pt(size_pt)

    # Required for Word to respect the font fully
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)

# =======================

def set_document_font(doc: Document, name="Arial", size_pt=10):
    # normal text
    set_style_font(doc.styles["Normal"], name, size_pt)

    # headings
    for i in range(1, 10):
        means = f"Heading {i}"
        if means in doc.styles:
            set_style_font(doc.styles[means], name, size_pt + (4 if i == 1 else 2))

    # lists
    for style_name in [
        "List Bullet",
        "List Number",
        "List Bullet 2",
        "List Number 2",
    ]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], name, size_pt)

# =======================

def set_moderate_margins(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# =========================

# add summary text to a Word document

def add_summary_to_doc(
    doc: Document, 
    summary: str, 
) -> None:
    
    lines = [ln.rstrip() for ln in summary.splitlines(keepends=True)]

    for ln in lines:
        line = ln.strip()

        if not line:
            doc.add_paragraph("")
            continue

        add_md_line(doc, line)
        
    return None

# =======================

# write summaries to a Word document

def write_summaries_docx(
    records: List[Dict], 
    round_code: str,
    round_description: str,
    out_path: str, 
) -> None:

    doc = Document()
    
    set_document_font(doc, "Arial", 10)
    set_moderate_margins(doc)

    # global paragraph spacing
    normal = doc.styles["Normal"]
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15

    # headings
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        hpf = h.paragraph_format
        hpf.space_before = Pt(9)
        hpf.space_after = Pt(6)

    doc.add_heading("Candidate Summary Report", level=1)
    doc.add_heading(round_code, level=4)
    doc.add_heading(round_description, level=4)
    doc.add_paragraph("")

    #generated_at = time.strftime("%Y-%m-%d %H:%M")
    #doc.add_heading(f"Generated at {generated_at}", level=4)

    doc.add_paragraph("")
    doc.add_heading("Notes regarding journal/conference article lists", level=4)
    doc.add_paragraph("The articles listed for each candidate are the articles published in journals/conference proceedings given the supplied journal list only.", style="List Bullet")
    doc.add_paragraph("The number of citations for an article is shown in square brackets.", style="List Bullet")
    doc.add_paragraph("Articles for a particular journal are ordered by the number of citations.", style="List Bullet")
    doc.add_page_break()

    for idx, record in enumerate(records, start=1):
        summary = record.get("summary_markdown", "") or ""
        if not summary.strip():
            continue
        
        # add this candidate's summary
        add_summary_to_doc(doc, summary)
        
        doc.add_page_break()

    doc.save(out_path)
    return None

# =========================

# get basic profile info from HR spreadsheet 

def get_basic_candidate_info(candidate: tuple) -> Dict:
    
    return {
            "candidate_id": candidate.candidate_id,
            "candidate_name": candidate.candidate_name,
            "gender": candidate.gender,
            "email": candidate.email,
            "country": candidate.country,
            "current_employee": candidate.current_employee,
            "expertise_area": candidate.expertise_area,
            "academic_level": candidate.academic_level,
            "PhD_year": candidate.PhD_year,
            "gs_url": candidate.gs_url,  
            "PhD_institution": candidate.PhD_institution,
            "PhD_institution_rank": int(candidate.PhD_institution_rank) if str(candidate.PhD_institution_rank).isdigit() else float('nan'),
            "YNM": "",
            "comments": "",
            "recruiter_notes": "",        
        }

# =========================

# main driver to process a profile

def process_profile(
    candidate: tuple,
    journal_list: List[str],
    normalised_journal_titles: Dict[str, str],
    html_dir: str,
) -> Dict[str, object] | None:

    global OFFLINE_MODE
    info_found = False
    
    print(f" === Processing profile for candidate {candidate.candidate_id}: {candidate.candidate_name} ===\n")
        
    url = candidate.gs_url
                
    if pd.isna(url):
        print(" Warning - Empty Google Scholar Link, skipping profile.")
        return empty_record(candidate, journal_list)
            
    url = sanitise_url(str(url).strip())
    if url is None:
        print(" Warning - Could not sanitise URL, skipping profile.")
        return empty_record(candidate, journal_list)
    
    record = get_basic_candidate_info(candidate)
    
    try:
        (
            gs_name,
            gs_institution,
            gs_research_areas,
            h_all,
            h_5y,
            cit_all,
            cit_5y,
            journal_counts,
            journal_counts_fa,
            journal_counts_sa,
            journal_counts_la,
            journal_num_authors,
            journal_details,
            article_count,
            info_found,
        ) = scrape_profile_all_publications(
            profile_url=url,
            journal_list=journal_list,
            normalised_journal_titles=normalised_journal_titles,
            html_dir=html_dir,
        )
        
    except AuthorMatchError:
        raise        
    except Exception as e:
        print(f" ERROR  - Detected error while processing cached HTML for {url}.")
        print(f" Details: {e}")
        return empty_record(candidate, journal_list)

    if not info_found:
        print(f" Warning - No scrapable pages found for URL: {url}")
        return empty_record(candidate, journal_list)

    if (
        gs_name is None
        and gs_institution is None
        and not gs_research_areas
        and all(v == 0 for v in journal_counts.values())
    ):
        print(f" Warning - No meaningful data scraped for URL: {url}")
        return empty_record(candidate, journal_list)
    
    average_num_authors = (
        round(
            sum(journal_num_authors.values()) / sum(journal_counts.values()),
            1
        )
        if sum(journal_counts.values()) > 0
        else 0.0
    )
    
    record.update({    
        "gs_name": gs_name or "",
        "gs_institution": gs_institution or "",
        "gs_research_areas": "; ".join(gs_research_areas) if gs_research_areas else "",
        "citations_all": cit_all if cit_all is not None else "",
        "citations_5y": cit_5y if cit_5y is not None else "",
        "h_index_all": h_all if h_all is not None else "",
        "h_index_5y": h_5y if h_5y is not None else "",
        "article_count": article_count if article_count is not None else "",
        "journal_count_tot": sum(journal_counts.values()),
        "journal_count_tot_fa": sum(journal_counts_fa.values()),
        "journal_count_tot_sa": sum(journal_counts_sa.values()),
        "journal_count_tot_la": sum(journal_counts_la.values()),
        "journal_average_num_authors": average_num_authors,
    })

    record["summary_markdown"] = create_summary(
        record=record,
        journal_counts=journal_counts,
        journal_counts_fa=journal_counts_fa,
        journal_counts_sa=journal_counts_sa,
        journal_counts_la=journal_counts_la,
        journal_num_authors=journal_num_authors,        
        journal_details=journal_details,
        journal_list=journal_list,
        is_empty_record=False,
        markdown=True
    )
    
    record["summary_plaintext"] = create_summary(
        record=record,
        journal_counts=journal_counts,
        journal_counts_fa=journal_counts_fa,
        journal_counts_sa=journal_counts_sa,
        journal_counts_la=journal_counts_la,
        journal_num_authors=journal_num_authors,        
        journal_details=journal_details,
        journal_list=journal_list,
        is_empty_record=False,
        markdown=False
    )
        
    for j in journal_list:
        record[j] = journal_counts.get(j, 0)

    return record

# =========================

# return an empty record for failed profiles

def empty_record(
    candidate: tuple,
    journal_list: List[str]
) -> Dict[str, object]:
    
    global NOT_FOUND_STRING
    global NOT_FOUND_NAN
   
    print(f"\n Setting empty record \n")
    print("\n ===============================================================================\n")
    record = get_basic_candidate_info(candidate=candidate)
    
    record.update({   
        "gs_name": NOT_FOUND_STRING,
        "gs_institution": NOT_FOUND_STRING,
        "gs_research_areas": NOT_FOUND_STRING,
        "citations_all": NOT_FOUND_NAN,
        "citations_5y": NOT_FOUND_NAN,
        "h_index_all": NOT_FOUND_NAN,
        "h_index_5y": NOT_FOUND_NAN,
        "article_count": NOT_FOUND_NAN,
        "journal_count_tot": NOT_FOUND_NAN,
        "journal_count_tot_fa": NOT_FOUND_NAN,
        "journal_count_tot_sa": NOT_FOUND_NAN,
        "journal_count_tot_la": NOT_FOUND_NAN,
        "journal_average_num_authors": NOT_FOUND_NAN,
    })

    record["summary_markdown"] = create_summary(
        record=record,
        is_empty_record=True,
        markdown=True
    )
    
    record["summary_plaintext"] = create_summary(
        record=record,
        is_empty_record=True,
        markdown=False
    )    
   
    for j in journal_list:
        record[j] = NOT_FOUND_NAN
    return record

# =========================

# open default web browser to a URL

def open_default_browser(url: str = "https://www.google.com") -> bool:
    try:
        opened = webbrowser.open(url, new=2)
        if opened:
            print(f" Opened default browser with URL: {url}")
            return True
        else:
            print(f" Warning - Browser did not open URL: {url}")
            return False
    except Exception as e:
        print(" ERROR - Failed to open the default web browser.")
        print(f" Exception: {type(e).__name__}: {e}")
        return False

# =========================
# main
# =========================

def main():
    
    global OFFLINE_MODE
    global FETCH_ONLY_MODE
    global NORMAL_MODE
    global BLOCKING_SUSPECTED
    global DEBUG_MODE
    
    parser = argparse.ArgumentParser(
        description="Snappy - Super Neat Academic Profile Parser"
    )
    mode_group = parser.add_mutually_exclusive_group()
    mode_group.add_argument(
        "--offline",
        action="store_true",
        help="Only parse existing cached HTML in user/html; do not fetch from Google Scholar.",
    )
    mode_group.add_argument(
        "--fetch-only",
        action="store_true",
        help="Only fetch and cache HTML from Google Scholar; do not parse or write outputs.",
    )
    mode_group.add_argument(
        "--normal",
        action="store_true",
        help="Normal mode - fetch, parse, and write outputs (default).",
    )
    
    parser.add_argument(
        "--debug",
        action="store_true",
        help="Debug mode - gives verbose outputs.",
    )
    
    parser.add_argument(
        "--accept-defaults",
        action="store_true",
        help="Accept all default settings and run.",
    )


    args = parser.parse_args()
    OFFLINE_MODE = args.offline
    FETCH_ONLY_MODE = args.fetch_only
    NORMAL_MODE = args.normal   
    DEBUG_MODE = args.debug
    ACCEPT_DEFAULTS = args.accept_defaults
    
    print("\n ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")
    print(" ~~~~~~ Welcome to Snappy - The Super Neat Academic Profile Parser.py ~~~~~~")
    print(" ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~\n")

    # get relative path for input/output files
    cwd = os.getcwd()
    print(f" Current working directory: {cwd}")
    if cwd.endswith("snappy/src") or cwd.endswith("snappy\\src"):
        rel_path = "../user/"
    elif cwd.endswith("snappy"):
        rel_path = "./user/"
    elif cwd.endswith("snappy/user") or cwd.endswith("snappy\\user"):
        rel_path = "./"
    else:
        print(f"\n ERROR - Please run this script from within the 'snappy/user' directory.\n")
        return
    
    # give user a chance to invoke these modes if not already specified
    if not OFFLINE_MODE and not FETCH_ONLY_MODE:
        # give user option to run in offline mode
        if ACCEPT_DEFAULTS:
            print("\n Accepting default settings: running in NORMAL mode.\n")
        else:
            answer = input(
                "\n Do you want to run in OFFLINE mode (parse cached HTML only)? (y/N): "
            ).strip().lower()
            if answer == "y":
                OFFLINE_MODE = True

            if not OFFLINE_MODE:
                # give user option to run in fetch-only mode
                answer = input(
                    "\n Do you want to run in FETCH-ONLY mode (download and cache pages only)? (y/N): "
                ).strip().lower()
                if answer == "y":
                    FETCH_ONLY_MODE = True
                
    if OFFLINE_MODE:
        print("\n Running in OFFLINE mode (shhh!): I will not contact Google Scholar.")
        print(" Instead I will parse HTML files already present in the 'html' directory.\n")
    elif FETCH_ONLY_MODE:
        print("\n Running in FETCH-ONLY mode: I will download and cache pages from Google Scholar, ")
        print(" but will not parse or write to output files.\n")
    else:
        print("\n Running in NORMAL mode: I will download and parse Google Scholar profile pages.\n")
                        
    # request HR report name
    if ACCEPT_DEFAULTS:
        print(
            " Accepting default HR report file name: "
            "'Campaign_Application_Report.xlsx'.\n"
        )
        hr_report_file = "Campaign_Application_Report.xlsx"
    else:
        hr_report_file = input(
            " Enter the name of the HR report file to process "
            "or press Enter for default ('Campaign_Application_Report.xlsx'):\n "
        ) or "Campaign_Application_Report.xlsx"
    
    hr_report_file = rel_path + hr_report_file.strip()
    
    if not os.path.exists(hr_report_file):
        print(f" ERROR - {hr_report_file} not found in current directory.")
        return    
    
    # convert the report to CSV for easier processing
    print(f"\n Extracting information from '{hr_report_file}' for processing...")
    try:
        df_hr = pd.read_excel(hr_report_file, header=None)
        
        # read the first row
        round_description = df_hr.iloc[0][0].strip()
        print (f" Detected application full description: {round_description}\n")
        round_code = round_description.split(" ")[0].strip()
        print (f" Using application round code: {round_code}\n")
        # remove round_code from round_description
        round_description = round_description[len(round_code):].strip(" -") 
        # remove "- IN CONFERENCE" if present
        round_description = round_description.replace(" - IN CONFERENCE", "").strip()
        
        # remove the first two rows of the data frame
        df_hr = df_hr.iloc[2:].reset_index(drop=True)

        # take the next row as header, clean its newlines, then set as columns
        raw_header = df_hr.iloc[0]

        # convert to string and strip/replace newlines
        clean_header = (
            raw_header
            .astype(str)
            .str.replace(r"[\r\n]+", " ", regex=True)
            .str.strip()
        )
        
        # print the header
        print(" Detected HR report columns:")
        for i, col in enumerate(clean_header):
            print(f"  {i + 1:02d}. {col}")

        # use this cleaned row as the header
        df_hr = df_hr[1:].reset_index(drop=True)
        df_hr.columns = clean_header

        # also: just in case, clean any lingering newlines in column names
        df_hr.columns = [
            str(c).replace("\r", " ").replace("\n", " ").strip()
            for c in df_hr.columns
        ]

        # clean up string-like columns to remove newline chars from cell values
        for col in df_hr.columns:
            if is_string_dtype(df_hr[col]):
                df_hr[col] = df_hr[col].map(
                    lambda x: x.replace("\r", " ").replace("\n", " ")
                    if isinstance(x, str) else x
                )

    except Exception as e:
        print(f" ERROR - Could not convert HR report to CSV. Exception: {type(e).__name__}: {e}")
        return
        
    # rename columns to something manageable and without grammatical errors ... 
    df_hr = df_hr.rename(columns={
        "Candidate Name": "candidate_name",
        "Candidate": "candidate_id",
        "Gender": "gender",
        "Email Address": "email",
        "In what country do you currently reside in?": "country",
        "Are you a student or current employee?": "current_employee",
        "What is your area of expertise?": "expertise_area",
        "What is the Academic Level you are applying for?": "academic_level",
        "Which year did you obtain your PhD? (YYYY)(Required if you have completed a PhD)": "PhD_year",
        "Which Institution did you obtain your PhD from?": "PhD_institution",
        "PhD Institution Rank": "PhD_institution_rank",
        "Google Scholar Link": "gs_url",
        "Would you like to longlist/Shortlist this candidate? Y= Yes M = Maybe N =No": "YNM",
        "Comments": "comments",
        "Recruiter Notes": "recruiter_notes",
    })
   
    print("\n ------------------------------------------\n")

    # read list of key journal names
    if ACCEPT_DEFAULTS:
        print(" Accepting default journal list file name: 'journal_list.txt'.\n")
        journal_list_file = "journal_list.txt"
    else:
        journal_list_file = input(
            " Enter the name of the file containing a list of journal names of interest "
            "or press Enter for default ('journal_list.txt'):\n "
        ) or "journal_list.txt"
    
    journal_list_file = rel_path + journal_list_file.strip()
    
    if not os.path.exists(journal_list_file):
        print(f" Warning - {journal_list_file} not found. I will not count journal publications.")
        journal_list: List[str] = []
    else:
        with open(journal_list_file, "r", encoding="utf-8") as f:
            journal_list = [line.strip() for line in f if line.strip()]

        if not journal_list:
            print(" Warning - No journal titles found. No journal counts will be recorded.")
        else:
            print(f" Loaded {len(journal_list)} journal titles from {journal_list_file}")
    
            # sort journal list in alphabetical order
            journal_list.sort()
            
            # remove duplicates from journal list
            journal_list = list(dict.fromkeys(journal_list))
            print(f" After removing duplicates, {len(journal_list)} unique journal titles will be used.\n")

        normalised_journal_titles = {
            normalise_journal_name(j): j
            for j in journal_list
        }
    
    # html caching
    html_dir = rel_path + "html"

    if OFFLINE_MODE:
        # in offline mode we never write HTML, we just read from html_dir
        os.makedirs(html_dir, exist_ok=True)
        print(f"\n Offline mode: I will use cached HTML pages under: {html_dir}")
    else:
        os.makedirs(html_dir, exist_ok=True)
        print(f" Normal or Fetch-only mode: I will cache HTML pages under: {html_dir}")

    # ask user to enter the candidate number to start on (default 1)
    if ACCEPT_DEFAULTS:
        print("\n Accepting default start candidate number: 1.\n")
        start_candidate_num = 1
    else:
        start_candidate_num_str = input("\n Enter the candidate number to start processing from (default 1):\n ")
        if not start_candidate_num_str.strip():
            start_candidate_num = 1
        else:
            try:
                start_candidate_num = int(start_candidate_num_str.strip())
            except ValueError:
                print(" Invalid candidate number entered, defaulting to 1.")
                start_candidate_num = 1
        
        if start_candidate_num < 1 or start_candidate_num > len(df_hr):
            print(f" Invalid candidate number {start_candidate_num}, must be between 1 and {len(df_hr)}. Defaulting to 1.")
            start_candidate_num = 1
        print(f" Starting processing from candidate number {start_candidate_num}...\n")
        df_hr = df_hr.iloc[start_candidate_num - 1 :].reset_index(drop=True)
    
    # ask user to enter the candidate number to stop on (default last)
    default_last = 5## len(df_hr) + start_candidate_num - 1
    if ACCEPT_DEFAULTS:
        print(f"\n Accepting default last candidate number: {default_last}.\n")
        end_candidate_num = default_last
    else:
        end_candidate_num_str = input(
            f"\n Enter the candidate number to stop processing on (default {default_last}):\n "
        )
        if not end_candidate_num_str.strip():
            end_candidate_num = default_last
        else:
            try:
                end_candidate_num = int(end_candidate_num_str.strip())
            except ValueError:
                print(f" Invalid candidate number entered, defaulting to {default_last}.")
                end_candidate_num = default_last
        if end_candidate_num < start_candidate_num or end_candidate_num > (default_last):
            print(f" Invalid candidate number {end_candidate_num}, must be between {start_candidate_num} and {default_last}.")
            end_candidate_num = default_last
    print(f" Stopping processing on candidate number {end_candidate_num}...\n")
    df_hr = df_hr.iloc[: end_candidate_num - start_candidate_num + 1].reset_index(drop=True)
    
    if not OFFLINE_MODE:
        if ACCEPT_DEFAULTS:
            print("\n Accepting default delay and retry settings.\n")
            typical_delay = 8.0
            max_block_retries = MAX_BLOCK_RETRIES_DEFAULT
        else:
            # get the typical delay to use to avoid blocking
            print("\n To reduce the chance of being blocked by Google Scholar, I need to wait for a random time period between web requests.")
            print(" The longer the wait time, the less likely you are to be blocked, but the longer the total processing time will be.\n")
            typical_delay_str = input(" Enter the typical delay in seconds between requests (default 8.0):\n ") 
            if not typical_delay_str.strip():
                typical_delay = 8.0
            else:
                try:
                    typical_delay = float(typical_delay_str.strip())
                except ValueError:
                    print(" Invalid delay entered, defaulting to 8.0 seconds.")
                    typical_delay = 8.0
                    if typical_delay <= 0.0:
                        typical_delay = 8.0
                        print(" Sorry. Second law of thermodynamics forbids going backwards in time.")
            print(f" Using a typical delay of {typical_delay} seconds between requests.\n")
            
            # get max block retries
            max_block_retries_str = input(
                f" Enter the maximum number of retries if blocking is suspected (default {MAX_BLOCK_RETRIES_DEFAULT}):\n "
            )
            if not max_block_retries_str.strip():
                max_block_retries = MAX_BLOCK_RETRIES_DEFAULT
            elif max_block_retries_str.strip().isdigit():
                max_block_retries = int(max_block_retries_str.strip())
            else:
                print(f" Invalid input, defaulting to {MAX_BLOCK_RETRIES_DEFAULT}.")
                max_block_retries = MAX_BLOCK_RETRIES_DEFAULT
    else:
        typical_delay = 0.01  # minimal delay in offline mode
            
    # start fetching and caching
    print("\n Attempting to fetch and cache all web pages...\n")
    print("\n ===============================================================================\n")

    session: Optional[requests.Session] = None
    if not OFFLINE_MODE:
        session = requests.Session()

        for candidate in df_hr.itertuples(index=False):
            
            try:
                # fetch and cache the pages
                if fetch_and_cache_profile(
                    candidate=candidate,
                    session=session,
                    pagesize=100,
                    max_pages=50,
                    delay=typical_delay,
                    max_block_retries=max_block_retries,
                    html_dir=html_dir,
                ):
                    print(f"\n Successfully fetched and cached pages for candidate {candidate.candidate_id}.\n")
                    print(" ================================================================\n")
                    random_sleep(typical_delay=typical_delay)
                
                else: 
                    print(f"\n Unable to fetch pages for candidate {candidate.candidate_id}.\n")
                    print(" ================================================================\n")
                    
            except GSBlockedError as e:
                BLOCKING_SUSPECTED = True
                # give user the option to continue or stop
                print(f"\n I suspect that Google is blocking web requests. Would you like to continue or stop?")
                print(f" Note that, if you stop now, you can restart from this candidate number next time.\n Then do a final run in OFFLINE mode to capture all candidates in the spreadsheet.\n")
                answer = input(" Enter 'c' to continue, 's' to stop processing: ").strip().lower()
                if answer == "c":
                    print(f"\n Continuing processing... but if this happens again soon I strongly suggest you stop and come back later.\n")
                    BLOCKING_SUSPECTED = False
                    time.sleep(5.0)  # brief pause before continuing
                    continue
                else:
                    print(f"\n Stopping further processing due to suspected blocking by Google.")
                    print(f" Please try again in an hour or two.")
                    print(f"\n Bye!\n")
                    print(" ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~\n")
                    break

        if FETCH_ONLY_MODE:
            print("\n ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~\n") 
            print(" Fetch-only mode complete.")
            print(" Cached HTML files (if any) are in the 'html' directory.")
            print(" You can now rerun Snappy in OFFLINE mode to parse them without contacting Google Scholar.\n")
            print(" Bye!\n")       
            return
    
    # start scraping
    print("\n Now scraping the web pages for key research metrics...\n")
    print("\n ===============================================================================\n")
    records: List[Dict[str, object]] = []
    
    for candidate in df_hr.itertuples(index=False):
        record = process_profile(
            candidate=candidate,
            journal_list=journal_list,
            normalised_journal_titles=normalised_journal_titles,
            html_dir=html_dir,
        )
        if record is not None:
            records.append(record)
        else:
            print(f" Warning - No record returned for candidate {candidate.candidate_id}.")
            print(f"\n ===============================================================================\n")


    if not records:
        print(" No records to write (no profiles scraped). Bye!\n")
        return
    
    # write to output files    
    timestamp = time.strftime("%Y-%m-%d_%H-%M")
    
    fieldnames = list(records[0].keys())
    
    column_labels: Dict[str, str] = {
        # HR fields
        "candidate_id": "Candidate ID",
        "candidate_name": "Candidate Name",
        "gender": "Gender",
        "email": "Email Address",
        "country": "Country of Residence",
        "current_employee": "Current Employee / Student",
        "expertise_area": "Expertise Area",
        "academic_level": "Academic Level Applied For",
        "PhD_year": "PhD Completion Year",
        "gs_url": "Google Scholar URL",        # this is the sanitised URL
        "PhD_institution": "PhD Institution",
        "PhD_institution_rank": "PhD Institution Rank",
        "YNM": "Longlist/Shortlist (Yes/No/Maybe)",
        "comments": "Comments",
        "recruiter_notes": "Recruiter Notes",
        # GS fields
        "gs_name": "Google Scholar Name",
        "gs_institution": "Google Scholar Institution",
        "gs_research_areas": "Google Scholar Research Areas",
        "citations_all": "Citations (All)",
        "citations_5y": "Citations (5y)",
        "h_index_all": "H-index (All)",
        "h_index_5y": "H-index (5y)",
        "article_count": "Total Number of Publications",
        "journal_count_tot": "Total Number of Publications in Journal List",
        "journal_count_tot_fa": "Number of First Author Publications in Journal List",
        "journal_count_tot_sa": "Number of Second Author Publications in Journal List",
        "journal_count_tot_la": "Number of Last Author Publications in Journal List",
        "journal_average_num_authors": "Average Number of Authors in Journal List Publications",      
        "summary_markdown": "Full Candidate Research Summary - MD",
        "summary_plaintext": "Full Candidate Research Summary - Double Click to Open / Use Up & Down Cursor Arrows to Navigate within Field / Escape to Exit",        
    }

    for j in journal_list:
        column_labels[j] = f"{j}"

    pretty_headers = [column_labels[col] for col in fieldnames]

    # write to xlsx
    xlsx_output_file = rel_path + "snappy_report_" + round_code + "_" + timestamp + ".xlsx"
    
    print(f"\n Writing records to: {xlsx_output_file} ...")
    
    try:
        df = pd.DataFrame(records, columns=fieldnames)
        df.columns = pretty_headers
        # remove columns we don't want in the Excel output for now
        df = df.drop(columns=["Average Number of Authors in Journal List Publications"])
        df = df.drop(columns=["Full Candidate Research Summary - MD"])        
        df.to_excel(xlsx_output_file, index=False)    
    except Exception as e:
        print(f"\n ERROR - Could not write Excel file. Exception: {type(e).__name__}: {e}")
        return

    # write to docx
    docx_output_file = rel_path + "snappy_report_" + round_code + "_" + timestamp + ".docx"

    print(f"\n Writing summary report to: {docx_output_file} ...")
    try:
        write_summaries_docx(records, round_code, round_description, docx_output_file)
    except Exception as e:
        print(f"\n ERROR - Could not write Word document. Exception: {type(e).__name__}: {e}")
        return
        

    print("\n ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")
    print(f"\n All done! Wrote {len(records)} rows to {xlsx_output_file} and {docx_output_file}.")
    print("\n ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")

    # optional bonus step: step through URLs in default browser
    if not OFFLINE_MODE:
        print("\n Would you like to step through each of the URLs? (y/N): ", end="")
        choice = input().strip().lower()

        if choice == "y":
            print("\n Stepping through each URL in your default web browser...")
            urls = sanitise_urls(df_hr["gs_url"].dropna().astype(str).tolist())
            for url in urls:
                print(f"\n Opening URL: {url}")
                opened = open_default_browser(url)
                if not opened:
                    print(" Warning: Could not open browser. Stopping step-through.\n")
                    break
                time.sleep(0.2)
                input(" Press Enter to continue to the next URL or Ctrl-C to quit...")

    print(f"\n Bye!\n")

# =========================
# main entry point
# =========================

if __name__ == "__main__":
    try:
        main()
    except AuthorMatchError as e:
        print("\nFATAL ERROR:")
        print(e)
        print("\nAborting Snappy due to inconsistent author matching.\n")
        sys.exit(1)
